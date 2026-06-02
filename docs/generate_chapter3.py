"""
Generate Chapter 3: User Manual in Urdu as a Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── helpers ────────────────────────────────────────────────────────────────

def rtl(para):
    pPr = para._p.get_or_add_pPr()
    b = OxmlElement('w:bidi')
    pPr.insert(0, b)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def urdu_font(run, size=12):
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Jameel Noori Nastaleeq')
    # mark as complex script
    cs_size = OxmlElement('w:szCs')
    cs_size.set(qn('w:val'), str(int(size * 2)))
    rPr.append(cs_size)


def upara(doc, text, size=12, bold=False, sb=0, sa=6):
    p = doc.add_paragraph()
    rtl(p)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.bold = bold
    urdu_font(r, size)
    return p


def uheading(doc, text, level=1):
    p = doc.add_paragraph()
    rtl(p)
    sizes = {1: 20, 2: 15, 3: 13}
    colors = {1: (0x1F, 0x49, 0x7D), 2: (0x2E, 0x74, 0xB5), 3: (0x20, 0x63, 0x20)}
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 7)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    urdu_font(r, sizes[level])
    c = colors[level]
    r.font.color.rgb = RGBColor(*c)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBEBEB')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p


def ubullet(doc, text, size=11):
    p = doc.add_paragraph()
    rtl(p)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + text)
    urdu_font(r, size)
    return p


def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(8)
    return p


def flow_table(doc, steps):
    """Draw a simple vertical flowchart using a narrow table."""
    for i, step in enumerate(steps):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        cell = tbl.cell(0, 0)
        cell.width = Cm(10)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(step)
        cr.bold = True
        cr.font.name = 'Jameel Noori Nastaleeq'
        cr.font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0' if i % 2 == 0 else 'EAF4FB')
        tcPr.append(shd)
        if i < len(steps) - 1:
            ap = doc.add_paragraph('↓')
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def state_table(doc, rows):
    """Two-column table: state | meaning in Urdu."""
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = 'Table Grid'
    # header
    for ci, txt in enumerate(['حالت', 'مطلب']):
        cell = tbl.cell(0, ci)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(txt)
        cr.bold = True
        urdu_font(cr, 11)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
        cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, (state, meaning) in enumerate(rows):
        for ci, txt in enumerate([state, meaning]):
            cell = tbl.cell(ri + 1, ci)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            cr = cp.add_run(txt)
            if ci == 0:
                cr.font.name = 'Courier New'
                cr.font.size = Pt(9)
                cr.bold = True
            else:
                urdu_font(cr, 10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F8FD' if ri % 2 == 0 else 'FFFFFF')
            tcPr.append(shd)
    doc.add_paragraph()


def hw_table(doc, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = 'Table Grid'
    for ci, txt in enumerate(['اجزاء', 'تفصیل']):
        cell = tbl.cell(0, ci)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(txt)
        cr.bold = True
        urdu_font(cr, 11)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
        cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, (comp, detail) in enumerate(rows):
        for ci, txt in enumerate([comp, detail]):
            cell = tbl.cell(ri + 1, ci)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cr = cp.add_run(txt)
            urdu_font(cr, 10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EBF5FB' if ri % 2 == 0 else 'FFFFFF')
            tcPr.append(shd)
    doc.add_paragraph()


# ─── document ───────────────────────────────────────────────────────────────

doc = Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)

# ── Chapter title ────────────────────────────────────────────────────────────
p = doc.add_paragraph()
rtl(p)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run('باب سوم')
r.bold = True
urdu_font(r, 24)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
rtl(p2)
p2.paragraph_format.space_after = Pt(16)
r2 = p2.add_run('صارف رہنما')
r2.bold = True
urdu_font(r2, 20)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3.1  نظام کی تنصیب
# ════════════════════════════════════════════════════════════════════════════
uheading(doc, '3.1  نظام کی تنصیب', level=1)

upara(doc,
    'یہ سیکشن Adam نظام کو NVIDIA Jetson Orin Nano Super Developer Kit پر مرحلہ وار نصب کرنے کا طریقہ بیان کرتا ہے۔ '
    'نظام مکمل طور پر آف لائن چلتا ہے، مگر پہلی بار تنصیب کے وقت انٹرنیٹ کنیکشن کی ضرورت ہوتی ہے تاکہ ضروری '
    'پیکیجز اور ماڈل فائلیں ڈاؤن لوڈ ہو سکیں۔ ایک بار تنصیب مکمل ہونے کے بعد نظام بغیر کسی آن لائن سروس کے '
    'چلتا رہتا ہے۔',
    sa=8
)

uheading(doc, '3.1.1  ہارڈ ویئر کی ضروریات', level=2)

upara(doc,
    'نظام کو چلانے کے لیے درج ذیل ہارڈ ویئر کی ضرورت ہے:',
    sa=4
)

hw_table(doc, [
    ('NVIDIA Jetson Orin Nano Super Developer Kit', '8 GB میموری کے ساتھ — مرکزی پروسیسنگ یونٹ'),
    ('256 GB NVMe SSD', 'آپریٹنگ سسٹم، ماڈل فائلیں، دستاویزات اور ڈیٹا بیس'),
    ('Jetson Nano USB Audio Codec', 'دو بلٹ ان مائیکروفون کے ساتھ — آواز ریکارڈ کرنے کے لیے'),
    ('بیرونی اسپیکر', 'آواز میں جواب سنانے کے لیے'),
    ('ESP32-S3 (دو گول اسکرین)', 'نظام کی حالت ظاہر کرنے والی روبوٹک آنکھیں'),
    ('Ethernet یا Wi-Fi', 'صرف پہلی تنصیب کے وقت انٹرنیٹ کے لیے'),
])

uheading(doc, '3.1.2  سافٹ ویئر کی ضروریات', level=2)

upara(doc,
    'ہارڈ ویئر کے علاوہ درج ذیل سافٹ ویئر اور ماڈل فائلوں کی ضرورت ہے:',
    sa=4
)

for item in [
    'Ubuntu Server — JetPack / Jetson Linux کے ذریعے نصب کریں',
    'Python 3.12 یا اس سے نئے ورژن',
    'Docker اور Docker Compose',
    'Ollama — مقامی زبانی ماڈل چلانے کے لیے',
    'Vosk ماڈل فائل — آواز کی پہچان کے لیے (models/vosk/)',
    'Piper ماڈل فائل — متن سے آواز بنانے کے لیے (models/piper/)',
    'all-MiniLM-L6-v2 — مقامی embedding ماڈل (models/embeddings/)',
    'llama3.2:latest — Ollama کے ذریعے چلنے والا زبانی ماڈل',
]:
    ubullet(doc, item)

doc.add_paragraph()

uheading(doc, '3.1.3  تنصیب کے مراحل', level=2)

# Step 1
uheading(doc, 'مرحلہ اول — آپریٹنگ سسٹم کی تنصیب', level=3)
upara(doc,
    'NVIDIA کی سرکاری ویب سائٹ سے JetPack SDK ڈاؤن لوڈ کریں اور Jetson Orin Nano Super پر Ubuntu Server نصب کریں۔ '
    'تنصیب مکمل ہونے کے بعد سسٹم کو اپ ڈیٹ کریں:',
    sa=3
)
code(doc, 'sudo apt update && sudo apt upgrade -y')

# Step 2
uheading(doc, 'مرحلہ دوم — آڈیو لائبریریاں', level=3)
upara(doc,
    'USB آڈیو کوڈیک اور مائیکروفون کو درست طریقے سے چلانے کے لیے PortAudio اور ALSA لائبریریاں نصب کریں:',
    sa=3
)
code(doc, 'sudo apt install python3.12-dev portaudio19-dev libasound2-dev libportaudio2 alsa-utils')

# Step 3
uheading(doc, 'مرحلہ سوم — Docker کی تنصیب', level=3)
upara(doc,
    'Docker کو سرکاری Docker Ubuntu repository سے نصب کریں۔ تفصیلی ہدایات docs/docker_deployment.md میں موجود ہیں۔ '
    'تنصیب کے بعد Docker سروس فعال کریں:',
    sa=3
)
code(doc, 'sudo systemctl enable docker\nsudo systemctl start docker')

# Step 4
uheading(doc, 'مرحلہ چہارم — پروجیکٹ اور Python انوائرنمنٹ', level=3)
upara(doc,
    'پروجیکٹ کے فولڈر میں جائیں اور Python virtual environment بنائیں، پھر تمام ضروری لائبریریاں نصب کریں:',
    sa=3
)
code(doc, 'python3 -m venv .venv\n.venv/bin/python -m pip install -e ".[all]"')

# Step 5
uheading(doc, 'مرحلہ پنجم — Ollama اور زبانی ماڈل', level=3)
upara(doc,
    'Ollama کو اس کی سرکاری ویب سائٹ سے Linux کے لیے نصب کریں۔ نصب کے بعد مطلوبہ ماڈل ڈاؤن لوڈ کریں:',
    sa=3
)
code(doc, 'ollama pull llama3.2:latest')

# Step 6
uheading(doc, 'مرحلہ ششم — Vosk اور Piper ماڈل فائلیں', level=3)
upara(doc,
    'Vosk کا آواز پہچاننے والا ماڈل ڈاؤن لوڈ کر کے models/vosk/ فولڈر میں رکھیں۔ '
    'اسی طرح Piper کا ماڈل models/piper/ فولڈر میں رکھیں۔ '
    'all-MiniLM-L6-v2 embedding ماڈل models/embeddings/ فولڈر میں رکھیں۔ '
    'یہ تینوں ماڈل نظام کے آف لائن چلنے کے لیے ضروری ہیں۔',
    sa=8
)

# Step 7
uheading(doc, 'مرحلہ ہفتم — Docker Image بنانا', level=3)
upara(doc,
    'تمام ماڈل اور ترتیبات مکمل ہونے کے بعد Docker image بنائیں:',
    sa=3
)
code(doc, 'docker compose build')
doc.add_paragraph()

rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3.2  نظام کی ترتیب
# ════════════════════════════════════════════════════════════════════════════
uheading(doc, '3.2  نظام کی ترتیب', level=1)

upara(doc,
    'نظام کی تمام ترتیب ایک .env فائل کے ذریعے کی جاتی ہے جو پروجیکٹ کے مرکزی فولڈر میں رکھی جاتی ہے۔ '
    'اس فائل میں تمام اہم متغیرات کی قدریں لکھی جاتی ہیں اور Docker Compose اسے خود بخود پڑھتا ہے۔ '
    '.env.example فائل کو بطور نمونہ استعمال کریں اور اسے کاپی کر کے .env نام دیں۔',
    sa=8
)

uheading(doc, '3.2.1  بنیادی ترتیب', level=2)
upara(doc, 'ضروری متغیرات جو تمام طریقوں میں مشترک ہیں:', sa=3)
code(doc,
    'ASSISTANT_OLLAMA_MODEL=llama3.2:latest\n'
    'ASSISTANT_OLLAMA_URL=http://127.0.0.1:11434\n'
    'ASSISTANT_CHROMA_PATH=./data/chroma\n'
    'ASSISTANT_SQLITE_PATH=./data/assistant.db\n'
    'ASSISTANT_EMBEDDING_MODEL=./models/embeddings/all-MiniLM-L6-v2\n'
    'ASSISTANT_EMBEDDING_DEVICE=cpu\n'
    'ASSISTANT_PIPER_MODEL=./models/piper/voice.onnx\n'
    'ASSISTANT_AUDIO_DEVICE=default\n'
    'ASSISTANT_VOSK_MODEL=./models/vosk'
)

uheading(doc, '3.2.2  ڈیش بورڈ کی ترتیب', level=2)
upara(doc,
    'ڈیش بورڈ کو لوکل نیٹ ورک پر دستیاب کرانے کے لیے host کو 0.0.0.0 پر سیٹ کریں:',
    sa=3
)
code(doc,
    'ASSISTANT_DASHBOARD_HOST=0.0.0.0\n'
    'ASSISTANT_DASHBOARD_PORT=8092\n'
    'ASSISTANT_DASHBOARD_ANSWER_MODE=llm\n'
    'ASSISTANT_OLLAMA_NUM_PREDICT=120\n'
    'ASSISTANT_OLLAMA_NUM_CTX=1536'
)

uheading(doc, '3.2.3  آواز کی ترتیب', level=2)
upara(doc,
    'یہ متغیرات آواز کے مختلف پہلوؤں کو کنٹرول کرتے ہیں۔ wake word، گفتگو کا وقفہ، اور '
    'جواب کا طریقہ یہاں طے کیا جاتا ہے:',
    sa=3
)
code(doc,
    'ASSISTANT_WAKE_WORD_KEYWORD=adam\n'
    'ASSISTANT_WAKE_WORD_ALIASES=adam,addam,atom,add em,add him,at him\n'
    'ASSISTANT_CONTINUOUS_CONVERSATION=1\n'
    'ASSISTANT_FOLLOWUP_TIMEOUT_SECONDS=45\n'
    'ASSISTANT_ENABLE_WAKE_GREETING=1\n'
    'ASSISTANT_STT_SPEECH_START_TIMEOUT_SECONDS=10\n'
    'ASSISTANT_STT_PRE_SPEECH_MS=300\n'
    'ASSISTANT_ENABLE_BARGE_IN=1\n'
    'ASSISTANT_VOICE_ANSWER_MODE=fast\n'
    'ASSISTANT_VOICE_RETRIEVAL_LIMIT=2\n'
    'ASSISTANT_VOICE_CONTEXT_CHARS=900'
)
upara(doc,
    'ASSISTANT_VOICE_ANSWER_MODE=fast کم وقفے کا راستہ ہے۔ اگر جواب کو Ollama کے ذریعے مزید تفصیلی بنانا ہو '
    'تو اسے llm پر سیٹ کریں۔',
    sa=8
)

uheading(doc, '3.2.4  آڈیو ڈیوائس کی جانچ', level=2)
upara(doc,
    'ترتیب مکمل کرنے کے بعد مائیکروفون اور اسپیکر کو جانچیں:',
    sa=3
)
code(doc, 'arecord -l\naplay -l')
upara(doc,
    'USB آڈیو کوڈیک عموماً plughw:1,0 یا plughw:0,3 کے طور پر ظاہر ہوتا ہے۔ '
    'اسے ASSISTANT_AUDIO_DEVICE متغیر میں لکھیں۔ TTS کی جانچ کے لیے:',
    sa=3
)
code(doc, 'PYTHONPATH=src python3 -m assistant_app.speak "Hello, I am ready."')

uheading(doc, '3.2.5  ESP32-S3 آنکھوں کی ترتیب', level=2)
upara(doc,
    'ESP32-S3 کو USB کیبل سے Jetson سے جوڑیں اور serial port معلوم کریں:',
    sa=3
)
code(doc, 'ls /dev/ttyUSB*')
upara(doc, 'پھر .env فائل میں یہ سطریں شامل کریں:', sa=3)
code(doc,
    'ASSISTANT_EYE_SERIAL_PORT=/dev/ttyUSB0\n'
    'ASSISTANT_EYE_SERIAL_BAUDRATE=115200'
)
doc.add_paragraph()

rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3.3  نظام کا استعمال
# ════════════════════════════════════════════════════════════════════════════
uheading(doc, '3.3  نظام کا استعمال', level=1)

upara(doc,
    'Adam کو دو طریقوں سے استعمال کیا جا سکتا ہے: آواز کے ذریعے براہ راست بات کر کے، یا '
    'ویب ڈیش بورڈ پر سوال ٹائپ کر کے۔ دونوں طریقے ایک ہی مقامی RAG pipeline استعمال کرتے ہیں '
    'جو اپ لوڈ کردہ دستاویزات میں سے جواب تلاش کرتی ہے۔',
    sa=8
)

uheading(doc, '3.3.1  نظام شروع کرنا', level=2)
upara(doc, 'Docker کے ذریعے نظام شروع کریں:', sa=3)
code(doc,
    '# تمام ضروری تصاویر بنائیں\n'
    'docker compose build\n\n'
    '# دستاویزات کا index بنائیں\n'
    'docker compose --profile tools run --rm ingest\n\n'
    '# نظام چلائیں\n'
    'docker compose up'
)
upara(doc,
    'نظام شروع ہونے کے بعد ڈیش بورڈ پورٹ 8092 پر دستیاب ہو جاتا ہے اور آواز کا نظام '
    'wake word سننا شروع کر دیتا ہے۔',
    sa=8
)

uheading(doc, '3.3.2  دستاویزات اپ لوڈ اور انڈیکس کرنا', level=2)
upara(doc,
    'Adam جوابات صرف انہی دستاویزات سے دیتا ہے جو پہلے سے انڈیکس کی گئی ہوں۔ '
    'دستاویز اپ لوڈ کرنے کے دو طریقے ہیں:',
    sa=4
)

upara(doc, 'طریقہ اول — ویب ڈیش بورڈ کے ذریعے:', bold=True, sa=2)
for step in [
    'براؤزر میں https://[device-ip]:8092 کھولیں',
    'ہوم پیج پر "Upload Documents" بٹن پر کلک کریں',
    'PDF، TXT یا MD فائل منتخب کریں',
    'فائل اپ لوڈ ہونے کے بعد "Ingest Documents" بٹن دبائیں',
    'انڈیکس مکمل ہونے کا پیغام ملنے کے بعد سوال پوچھیں',
]:
    ubullet(doc, step)

doc.add_paragraph()
upara(doc, 'طریقہ دوم — کمانڈ لائن سے:', bold=True, sa=2)
upara(doc, 'دستاویز کو documents/ فولڈر میں رکھیں، پھر:', sa=3)
code(doc, 'assistant-ingest ./documents')
doc.add_paragraph()

uheading(doc, '3.3.3  آواز کا طریقہ استعمال', level=2)
upara(doc,
    'نظام شروع ہونے کے بعد آواز کے نظام کا مرحلہ وار طریقہ یہ ہے:',
    sa=4
)

flow_table(doc, [
    '① "Adam" کہیں — wake word',
    '② نظام کا مختصر جواب سنیں ("Yes?" یا "I\'m listening")',
    '③ اپنا سوال واضح آواز میں پوچھیں',
    '④ نظام سوچتا ہے اور "One moment..." جیسا فیڈبیک دیتا ہے',
    '⑤ دستاویزات میں تلاش کے بعد آواز میں جواب ملتا ہے',
    '⑥ 45 سیکنڈ تک آگے کا سوال بغیر wake word کے پوچھیں',
    '⑦ کوئی سوال نہ ہو تو نظام خود بخود wake word mode میں واپس آ جاتا ہے',
])

upara(doc,
    'شور والے ماحول میں wake word کے متبادل الفاظ جیسے "atom" یا "add him" بھی کام کرتے ہیں '
    'کیونکہ یہ ASSISTANT_WAKE_WORD_ALIASES میں پہلے سے شامل ہیں۔',
    sa=8
)

uheading(doc, '3.3.4  سروس کا انتظام', level=2)
upara(doc,
    'اگر adam.service کے طور پر نصب ہے تو یہ کمانڈز استعمال کریں:',
    sa=3
)
code(doc,
    'sudo systemctl start adam.service\n'
    'sudo systemctl stop adam.service\n'
    'sudo systemctl restart adam.service\n'
    'systemctl status adam.service --no-pager\n'
    'journalctl -u adam.service --since "2 minutes ago" --no-pager'
)
doc.add_paragraph()

rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3.4  صارف انٹرفیس کی تفصیل
# ════════════════════════════════════════════════════════════════════════════
uheading(doc, '3.4  صارف انٹرفیس کی تفصیل', level=1)

upara(doc,
    'Adam کا صارف انٹرفیس تین حصوں پر مشتمل ہے: ویب ڈیش بورڈ، آواز کا انٹرفیس، '
    'اور ESP32-S3 کی روبوٹک آنکھیں۔ یہ تینوں مل کر ایک مکمل تجربہ فراہم کرتے ہیں۔',
    sa=8
)

uheading(doc, '3.4.1  ویب ڈیش بورڈ', level=2)
upara(doc,
    'ڈیش بورڈ HTTPS پر پورٹ 8092 پر چلتا ہے۔ براؤزر میں کھولنے کا پتہ:',
    sa=3
)
code(doc, 'https://[device-ip]:8092')

upara(doc, 'ڈیش بورڈ کے اہم صفحات:', sa=4, bold=True)

uheading(doc, '/ask — سوال پوچھنے کا صفحہ', level=3)
upara(doc,
    'یہ نظام کا مرکزی چیٹ انٹرفیس ہے جہاں صارف لکھ کر سوال پوچھتا ہے۔ '
    'اس صفحے پر درج ذیل سہولیات موجود ہیں:',
    sa=3
)
for feat in [
    'New Chat — نئی گفتگو شروع کرنے کا بٹن',
    'حالیہ گفتگو کی فہرست — پچھلے سیشنز تک فوری رسائی',
    'Streamed جواب — جواب حرف حرف آتا ہے، انتظار کم محسوس ہوتا ہے',
    'Source chips — جواب کس دستاویز اور صفحے سے ملا یہ ظاہر ہوتا ہے',
    'Copy — جواب کاپی کرنے کا بٹن',
    'سوال نہ ملنے پر عمومی receptionist جواب ملتا ہے',
]:
    ubullet(doc, feat)

doc.add_paragraph()

uheading(doc, '/monitor — نگرانی کا صفحہ', level=3)
upara(doc,
    'یہ صفحہ نظام کے منتظمین اور محققین کے لیے ہے۔ یہاں یہ معلومات ملتی ہیں:',
    sa=3
)
for item in [
    'index کی صحت اور دستاویزات کی تعداد',
    'چیٹ ریکارڈز اور آواز کے واقعات',
    'جواب کے اوسط وقت (latency) کی اعداد و شمار',
    'Ollama اور ماڈل کی حالت',
    'سسٹم کے وسائل (CPU، میموری)',
    'غلطیوں کا ریکارڈ',
    'تمام ریکارڈز CSV/JSON میں ایکسپورٹ کرنے کے لنک',
]:
    ubullet(doc, item)

doc.add_paragraph()
upara(doc,
    '/monitor/export/interactions.csv — تمام سوال جوابات ایکسپورٹ کریں۔\n'
    '/monitor/interaction/{id} — ایک مخصوص گفتگو کی مکمل تفصیل دیکھیں۔',
    sa=8
)

uheading(doc, '3.4.2  آواز کا انٹرفیس', level=2)
upara(doc,
    'آواز کا انٹرفیس مکمل طور پر قدرتی زبان پر مبنی ہے۔ '
    'صارف صرف "Adam" کہہ کر نظام سے بات کرتا ہے، کوئی بٹن نہیں دبانا پڑتا۔ '
    'نظام مختلف حالتوں میں مختلف آوازیں دیتا ہے جیسے تصدیقی آوازیں، سوچنے کی آوازیں، '
    'اور مکمل جواب۔',
    sa=8
)

uheading(doc, '3.4.3  ESP32-S3 روبوٹک آنکھیں', level=2)
upara(doc,
    'دو گول اسکرینوں والی ESP32-S3 ڈیوائس نظام کی موجودہ حالت بصری طور پر ظاہر کرتی ہے۔ '
    'یہ صارف کو بتاتی ہے کہ نظام ابھی کیا کر رہا ہے بغیر کوئی سوال کیے:',
    sa=4
)

state_table(doc, [
    ('IDLE',      'نظام تیار ہے، wake word کا انتظار'),
    ('LISTENING', 'آواز سن رہا ہے'),
    ('THINKING',  'جواب تیار کر رہا ہے'),
    ('SEARCHING', 'دستاویزات میں تلاش جاری ہے'),
    ('SPEAKING',  'جواب بول رہا ہے'),
    ('READY',     'جواب مکمل، اگلے سوال کا انتظار'),
    ('ERROR',     'کوئی خرابی پیش آئی'),
])

doc.add_paragraph()
rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3.5  خصوصیات اور افعال
# ════════════════════════════════════════════════════════════════════════════
uheading(doc, '3.5  خصوصیات اور افعال', level=1)

upara(doc,
    'Adam نظام کئی اہم خصوصیات کا مجموعہ ہے جو اسے ایک عام chatbot سے مختلف بناتی ہیں۔ '
    'یہ تمام خصوصیات مقامی طور پر اور بغیر کسی آن لائن سروس کے کام کرتی ہیں:',
    sa=8
)

uheading(doc, '3.5.1  مکمل آف لائن آپریشن', level=2)
upara(doc,
    'نظام کی سب سے اہم خصوصیت اس کا مکمل طور پر آف لائن چلنا ہے۔ '
    'صارف کی آواز، سوال، اور اپ لوڈ کردہ دستاویزات کا کوئی حصہ بھی کسی بیرونی سرور یا '
    'کلاؤڈ سروس کو نہیں بھیجا جاتا۔ یہ ادارہ جاتی رازداری کے لیے ضروری ہے خاص طور پر '
    'جب حساس دستاویزات انڈیکس کی جائیں۔',
    sa=8
)

uheading(doc, '3.5.2  آواز کی پہچان — Vosk', level=2)
upara(doc,
    'Vosk ایک مفت اور کھلے ذریعہ کوڈ کا speech-to-text انجن ہے جو مقامی طور پر چلتا ہے۔ '
    'نظام مسلسل مائیکروفون سے آواز لیتا رہتا ہے اور Vosk اسے متن میں بدلتا ہے۔ '
    'Voice Activity Detection (VAD) یہ طے کرتا ہے کہ صارف کب بول رہا ہے اور کب خاموش ہے۔ '
    'ایک چھوٹا pre-speech buffer سوال کے شروع کو محفوظ رکھتا ہے۔',
    sa=8
)

uheading(doc, '3.5.3  دستاویز تلاش — RAG Pipeline', level=2)
upara(doc,
    'Retrieval-Augmented Generation (RAG) نظام کا مرکزی دماغ ہے۔ '
    'جب صارف سوال پوچھتا ہے تو یہ pipeline ان مراحل سے گزرتی ہے:',
    sa=4
)

flow_table(doc, [
    'سوال کو embedding میں بدلنا — all-MiniLM-L6-v2',
    'ChromaDB میں semantic search — سب سے قریبی chunks تلاش',
    'SQLite سے metadata حاصل کرنا — صفحہ نمبر، ماخذ',
    'متعلقہ text chunks کو Ollama کو بھیجنا',
    'Ollama کا مقامی جواب تیار کرنا — llama3.2',
    'جواب کی تصدیق — دستاویز سے میل کھاتا ہے؟',
    'آواز یا ویب پر جواب دینا',
])

upara(doc,
    'اگر دستاویزات میں سوال کا جواب نہ ملے تو نظام عمومی receptionist انداز میں جواب دیتا ہے '
    'یا صارف سے متعلقہ دستاویز اپ لوڈ کرنے کی درخواست کرتا ہے۔',
    sa=8
)

uheading(doc, '3.5.4  متعدد موضوعاتی گفتگو', level=2)
upara(doc,
    'wake word کہنے کے بعد صارف 45 سیکنڈ تک آگے کے سوال بغیر "Adam" دہرائے پوچھ سکتا ہے۔ '
    'یہ گفتگو کو قدرتی اور روانی بخشتی ہے۔ وقفہ ختم ہونے پر نظام خود بخود wake word mode میں '
    'واپس آ جاتا ہے۔ ASSISTANT_FOLLOWUP_TIMEOUT_SECONDS سے یہ وقفہ تبدیل کیا جا سکتا ہے۔',
    sa=8
)

uheading(doc, '3.5.5  متن سے آواز — Piper', level=2)
upara(doc,
    'Piper ایک تیز رفتار مقامی text-to-speech انجن ہے جو Jetson Orin Nano پر بھی بہترین '
    'کام کرتا ہے۔ نظام جواب تیار ہونے پر اسے Piper کو بھیجتا ہے جو عارضی WAV فائل بناتا ہے '
    'اور aplay سے اسپیکر پر چلاتا ہے۔',
    sa=8
)

uheading(doc, '3.5.6  ویب ڈیش بورڈ اور ریکارڈ', level=2)
upara(doc,
    'SQLite ڈیٹا بیس تمام گفتگو، آواز کے واقعات، غلطیوں، اور دستاویز metadata کو محفوظ رکھتا ہے۔ '
    'یہ ریکارڈ CSV یا JSON میں ایکسپورٹ کیے جا سکتے ہیں جو تحقیق اور جائزے کے لیے مفید ہے۔',
    sa=8
)

uheading(doc, '3.5.7  بصری فیڈبیک — روبوٹک آنکھیں', level=2)
upara(doc,
    'ESP32-S3 کی آنکھیں نظام کو ایک شخصیت دیتی ہیں اور صارف کو ہر وقت یہ معلوم رہتا ہے '
    'کہ نظام کس مرحلے پر ہے۔ یہ صارف کے اعتماد اور تجربے کو بہتر بناتی ہیں۔',
    sa=8
)

uheading(doc, '3.5.8  Barge-in سہولت', level=2)
upara(doc,
    'ASSISTANT_ENABLE_BARGE_IN=1 کے ساتھ صارف جواب کے دوران بھی بول سکتا ہے اور نظام '
    'موجودہ جواب روک کر نیا سوال سننا شروع کر دیتا ہے۔ یہ گفتگو کو مزید قدرتی بناتا ہے۔',
    sa=8
)

rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3.6  نظام کی حدود
# ════════════════════════════════════════════════════════════════════════════
uheading(doc, '3.6  نظام کی حدود', level=1)

upara(doc,
    'Adam نظام کئی اہم کاموں کے لیے موزوں ہے، مگر کچھ تکنیکی اور عملی حدود بھی ہیں '
    'جن کا صارف اور منتظم کو علم ہونا ضروری ہے:',
    sa=8
)

uheading(doc, '3.6.1  پہلی تنصیب کے دوران انٹرنیٹ', level=2)
upara(doc,
    'نظام آف لائن چلتا ہے مگر پہلی بار تنصیب کے وقت انٹرنیٹ کنیکشن لازمی ہے۔ '
    'Ubuntu پیکیجز، Docker image، Ollama، Vosk ماڈل، Piper ماڈل، اور embedding ماڈل '
    'سب ڈاؤن لوڈ کرنے کی ضرورت ہوتی ہے جو مجموعی طور پر کئی گیگا بائٹ بنتے ہیں۔ '
    'ایک بار ڈاؤن لوڈ کے بعد انٹرنیٹ کی ضرورت نہیں۔',
    sa=8
)

uheading(doc, '3.6.2  دستاویزات کی فارمیٹ', level=2)
upara(doc,
    'نظام صرف PDF، TXT، اور Markdown (MD) فارمیٹ کی فائلیں پڑھ سکتا ہے۔ '
    'Word دستاویزات (.docx)، Excel (.xlsx)، PowerPoint (.pptx)، تصاویر، یا آڈیو/ویڈیو '
    'فائلیں ابھی سپورٹ نہیں کی جاتیں۔ ایسی فائلوں کو پہلے PDF میں تبدیل کرنا ہوگا۔',
    sa=8
)

uheading(doc, '3.6.3  آواز کی پہچان کی زبان', level=2)
upara(doc,
    'Vosk کا موجودہ ماڈل صرف انگریزی آواز کو متن میں بدلتا ہے۔ اردو، عربی، یا دیگر زبانوں '
    'میں بولنے والے صارفین کو انگریزی میں سوال پوچھنا ہوگا۔ دوسری زبانوں کے Vosk ماڈل '
    'علیحدہ ڈاؤن لوڈ کر کے ASSISTANT_VOSK_MODEL کی ترتیب تبدیل کی جا سکتی ہے مگر '
    'تجربہ کی کارکردگی انگریزی ماڈل جیسی نہیں ہوگی۔',
    sa=8
)

uheading(doc, '3.6.4  جواب کا وقت', level=2)
upara(doc,
    'جواب کی رفتار ہارڈ ویئر کی کارکردگی پر منحصر ہے۔ Jetson Orin Nano Super پر '
    'ASSISTANT_VOICE_ANSWER_MODE=fast کے ساتھ مختصر سوالوں کا جواب عموماً 1-2 سیکنڈ '
    'میں ملتا ہے۔ llm موڈ میں Ollama کے ذریعے جواب 3-8 سیکنڈ لے سکتا ہے۔ '
    'بڑی اور پیچیدہ دستاویزات کو انڈیکس کرنے میں بھی زیادہ وقت لگتا ہے۔',
    sa=8
)

uheading(doc, '3.6.5  صرف لوکل نیٹ ورک پر دستیاب', level=2)
upara(doc,
    'ڈیش بورڈ صرف اسی نیٹ ورک پر دستیاب ہے جس سے Jetson جڑا ہو۔ '
    'انٹرنیٹ سے براہ راست رسائی ممکن نہیں اور نظام کو اس کے لیے ڈیزائن بھی نہیں کیا گیا '
    'کیونکہ رازداری کا بنیادی اصول یہی ہے کہ تمام ڈیٹا مقامی رہے۔',
    sa=8
)

uheading(doc, '3.6.6  ایک صارف ایک وقت میں', level=2)
upara(doc,
    'آواز کا نظام ایک وقت میں ایک ہی صارف کی آواز سن اور سمجھ سکتا ہے۔ '
    'ایک سے زیادہ افراد بیک وقت بولیں تو نظام الجھ سکتا ہے۔ ویب ڈیش بورڈ پر ایک سے '
    'زیادہ صارف بیک وقت ٹائپ کر کے سوال پوچھ سکتے ہیں مگر Ollama بھی ایک وقت میں ایک '
    'درخواست پر کام کرتا ہے۔',
    sa=8
)

uheading(doc, '3.6.7  دستاویز کا مواد محدود', level=2)
upara(doc,
    'نظام صرف انہی معلومات کا جواب دے سکتا ہے جو اپ لوڈ کردہ دستاویزات میں موجود ہوں۔ '
    'اگر کسی موضوع پر دستاویز نہ ہو تو نظام یا تو عمومی receptionist انداز میں جواب دے گا '
    'یا صارف کو متعلقہ دستاویز اپ لوڈ کرنے کی گزارش کرے گا۔ نظام کسی بات کا غلط اندازہ '
    'لگانے کے بجائے ماخذ نہ ہونے کی اطلاع دینے کو ترجیح دیتا ہے۔',
    sa=8
)

rule(doc)

# ── save ─────────────────────────────────────────────────────────────────────
out = '/home/ubuntu/Adam/docs/Chapter_3_User_Manual.docx'
doc.save(out)
print(f'Saved: {out}')
